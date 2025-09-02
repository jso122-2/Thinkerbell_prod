"""
Document ingestion module for PDF, DOCX, and TXT files
"""

import logging
from pathlib import Path
from typing import Dict, List, Optional, Union
import mimetypes

# PDF parsing
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    fitz = None

# DOCX parsing
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    DocxDocument = None

# Text utilities
import re

from .utils import get_file_info, hash_id


logger = logging.getLogger(__name__)


class DocumentIngester:
    """
    Handles ingestion of various document formats (PDF, DOCX, TXT).
    
    Extracts text content and basic structure information from documents.
    """
    
    SUPPORTED_EXTENSIONS = {'.txt', '.md'}
    
    def __init__(self):
        """Initialize the document ingester."""
        # Add PDF support if available
        if PYMUPDF_AVAILABLE:
            self.SUPPORTED_EXTENSIONS.add('.pdf')
        
        # Add DOCX support if available  
        if DOCX_AVAILABLE:
            self.SUPPORTED_EXTENSIONS.add('.docx')
            
        self.stats = {
            'processed': 0,
            'errors': 0,
            'by_type': {}
        }
    
    def ingest_file(self, file_path: Union[str, Path]) -> Dict[str, any]:
        """
        Ingest a single document file.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing extracted text and metadata
            
        Raises:
            ValueError: If file type is not supported
            FileNotFoundError: If file doesn't exist
        """
        path_obj = Path(file_path)
        
        if not path_obj.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = path_obj.suffix.lower()
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {extension}")
        
        try:
            # Get file metadata
            source_path, mtime = get_file_info(path_obj)
            doc_id = hash_id(source_path, mtime)
            
            # Extract text based on file type
            if extension == '.pdf':
                text_content = self._extract_pdf(path_obj)
            elif extension == '.docx':
                text_content = self._extract_docx(path_obj)
            elif extension in ['.txt', '.md']:
                text_content = self._extract_text(path_obj)
            else:
                raise ValueError(f"Handler not implemented for {extension}")
            
            # Update statistics
            self.stats['processed'] += 1
            self.stats['by_type'][extension] = self.stats['by_type'].get(extension, 0) + 1
            
            result = {
                'source_path': source_path,
                'doc_id': doc_id,
                'extension': extension,
                'file_size': path_obj.stat().st_size,
                'mtime': mtime,
                'text_content': text_content,
                'sections': self._detect_sections(text_content)
            }
            
            logger.info(f"Successfully ingested {path_obj.name} ({len(text_content)} chars)")
            return result
            
        except Exception as e:
            self.stats['errors'] += 1
            logger.error(f"Error ingesting {file_path}: {str(e)}")
            raise
    
    def _extract_pdf(self, path: Path) -> str:
        """Extract text from PDF file using PyMuPDF."""
        if not PYMUPDF_AVAILABLE:
            raise RuntimeError("PyMuPDF is not installed. Install with: pip install PyMuPDF")
            
        try:
            doc = fitz.open(str(path))
            text_parts = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()
                if text.strip():
                    text_parts.append(text)
            
            doc.close()
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            raise RuntimeError(f"Failed to extract PDF text: {str(e)}")
    
    def _extract_docx(self, path: Path) -> str:
        """Extract text from DOCX file using python-docx."""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx is not installed. Install with: pip install python-docx")
            
        try:
            doc = DocxDocument(str(path))
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            raise RuntimeError(f"Failed to extract DOCX text: {str(e)}")
    
    def _extract_text(self, path: Path) -> str:
        """Extract text from plain text files."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            
            raise RuntimeError("Could not decode text file with any supported encoding")
            
        except Exception as e:
            raise RuntimeError(f"Failed to extract text: {str(e)}")
    
    def _detect_sections(self, text: str) -> Dict[str, str]:
        """
        Basic section detection for documents.
        
        Args:
            text: Full document text
            
        Returns:
            Dictionary with detected sections
        """
        sections = {
            'full': text,
            'header': '',
            'body': text,
            'footer': ''
        }
        
        lines = text.split('\n')
        if len(lines) < 3:
            return sections
        
        # Simple heuristic: first few lines might be header
        # Last few lines might be footer
        header_candidate = '\n'.join(lines[:3]).strip()
        footer_candidate = '\n'.join(lines[-3:]).strip()
        
        # Check for common header patterns
        header_patterns = [
            r'^\s*[A-Z\s]{10,}\s*$',  # All caps titles
            r'^\s*\d{1,2}[/-]\d{1,2}[/-]\d{2,4}',  # Dates
            r'^\s*Page\s+\d+',  # Page numbers
        ]
        
        for pattern in header_patterns:
            if re.search(pattern, header_candidate, re.MULTILINE):
                sections['header'] = header_candidate
                sections['body'] = '\n'.join(lines[3:]).strip()
                break
        
        # Check for footer patterns
        footer_patterns = [
            r'Page\s+\d+\s*of\s*\d+',
            r'©\s*\d{4}',  # Copyright
            r'Confidential|Proprietary',
        ]
        
        for pattern in footer_patterns:
            if re.search(pattern, footer_candidate, re.IGNORECASE):
                sections['footer'] = footer_candidate
                if sections['header']:
                    sections['body'] = '\n'.join(lines[3:-3]).strip()
                else:
                    sections['body'] = '\n'.join(lines[:-3]).strip()
                break
        
        return sections
    
    def ingest_directory(self, dir_path: Union[str, Path], recursive: bool = True) -> List[Dict[str, any]]:
        """
        Ingest all supported documents in a directory.
        
        Args:
            dir_path: Path to directory
            recursive: Whether to search subdirectories
            
        Returns:
            List of ingested document dictionaries
        """
        path_obj = Path(dir_path)
        if not path_obj.is_dir():
            raise ValueError(f"Not a directory: {dir_path}")
        
        results = []
        pattern = "**/*" if recursive else "*"
        
        for file_path in path_obj.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                try:
                    result = self.ingest_file(file_path)
                    results.append(result)
                except Exception as e:
                    logger.warning(f"Skipping {file_path}: {str(e)}")
                    continue
        
        logger.info(f"Ingested {len(results)} documents from {dir_path}")
        return results
    
    def get_stats(self) -> Dict[str, any]:
        """Get ingestion statistics."""
        return self.stats.copy()
    
    def reset_stats(self):
        """Reset ingestion statistics."""
        self.stats = {
            'processed': 0,
            'errors': 0,
            'by_type': {}
        } 