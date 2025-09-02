"""
Template ingestion and style profiling module.

Loads real PDF/DOCX templates, converts to clean text, and extracts
style profiles including tone, clause ordering, and phrasing patterns.
"""

import logging
import re
from pathlib import Path
from typing import Dict, List, Set, Tuple, Optional, Any
from collections import defaultdict, Counter
import json
import spacy

try:
    from tqdm import tqdm
    TQDM_AVAILABLE = True
except ImportError:
    TQDM_AVAILABLE = False
    # Fallback tqdm function
    def tqdm(iterable, *args, **kwargs):
        return iterable

# Enhanced imports for better text extraction
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    fitz = None

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    DocxDocument = None

from .ingest import DocumentIngester
from .clean import TextCleaner
from .utils import ensure_dir

logger = logging.getLogger(__name__)


class StyleProfile:
    """Represents a style profile extracted from real templates."""
    
    def __init__(self, profile_id: str):
        self.profile_id = profile_id
        self.tone_markers = {}  # formal/casual/professional indicators
        self.clause_patterns = []  # common clause structures
        self.phrase_vocabulary = set()  # key phrases and terminology
        self.sentence_structures = []  # typical sentence patterns
        self.industry_indicators = set()  # industry-specific terms
        self.formality_score = 0.0  # 0.0=casual, 1.0=very formal
        self.complexity_indicators = {}  # markers of document complexity
        self.template_count = 0  # number of templates in this profile
        
        # NEW: Required fields for eliminating unknown mappings
        self.industry = 'unknown'  # Will be determined from content
        self.complexity = 'unknown'  # Will be determined from content
        self.style_keywords = []  # Top N frequent nouns/adjectives
        
    def to_dict(self) -> Dict[str, Any]:
        return {
            'profile_id': self.profile_id,
            'tone_markers': self.tone_markers,
            'clause_patterns': self.clause_patterns,
            'phrase_vocabulary': list(self.phrase_vocabulary),
            'sentence_structures': self.sentence_structures,
            'industry_indicators': list(self.industry_indicators),
            'formality_score': self.formality_score,
            'complexity_indicators': self.complexity_indicators,
            'template_count': self.template_count,
            'industry': self.industry,
            'complexity': self.complexity,
            'style_keywords': self.style_keywords
        }
    
    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'StyleProfile':
        profile = cls(data['profile_id'])
        profile.tone_markers = data.get('tone_markers', {})
        profile.clause_patterns = data.get('clause_patterns', [])
        profile.phrase_vocabulary = set(data.get('phrase_vocabulary', []))
        profile.sentence_structures = data.get('sentence_structures', [])
        profile.industry_indicators = set(data.get('industry_indicators', []))
        profile.formality_score = data.get('formality_score', 0.0)
        profile.complexity_indicators = data.get('complexity_indicators', {})
        profile.template_count = data.get('template_count', 0)
        profile.industry = data.get('industry', 'unknown')
        profile.complexity = data.get('complexity', 'unknown')
        profile.style_keywords = data.get('style_keywords', [])
        return profile


class TemplateIngestor:
    """
    Ingests real templates and builds style profiles.
    
    Processes PDF/DOCX files to extract style patterns that can be used
    to influence synthetic data generation.
    """
    
    def __init__(self, templates_dir: Path, cache_dir: Optional[Path] = None):
        """
        Initialize template ingestor.
        
        Args:
            templates_dir: Directory containing template files
            cache_dir: Directory for caching (used for mapping diagnostics)
        """
        self.templates_dir = Path(templates_dir)
        self.cache_dir = Path(cache_dir) if cache_dir else Path("./synthetic_dataset/cache")
        self.ingester = DocumentIngester()
        self.cleaner = TextCleaner()
        
        # Initialize mapping diagnostics
        from .mapping_diagnostics import MappingDiagnostics
        self.mapping_diagnostics = MappingDiagnostics(self.cache_dir)
        
        # Initialize spaCy for better text analysis
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            logger.warning("spaCy model 'en_core_web_sm' not found. Using basic text analysis.")
            self.nlp = None
        
        # Style analysis patterns
        self._init_analysis_patterns()
        
        # Loaded templates and profiles
        self.templates = []
        self.style_profiles = {}
        
    def _init_analysis_patterns(self):
        """Initialize patterns for style analysis."""
        
        # Formality indicators
        self.formal_indicators = {
            'hereby', 'whereas', 'pursuant to', 'in accordance with',
            'notwithstanding', 'subject to', 'undersigned', 'hereinafter',
            'aforementioned', 'shall', 'covenant', 'warrant', 'represent'
        }
        
        self.casual_indicators = {
            'we\'ll', 'you\'ll', 'can\'t', 'won\'t', 'let\'s', 'hey',
            'awesome', 'cool', 'super', 'thanks', 'cheers'
        }
        
        # Enhanced Industry-specific terms
        self.industry_terms = {
            'entertainment': {
                'talent', 'agreement', 'performance', 'artistic', 'entertainment',
                'production', 'media', 'broadcast', 'influencer', 'social media',
                'content', 'brand', 'endorsement', 'promotional', 'campaign',
                'instagram', 'youtube', 'tiktok', 'facebook', 'twitter'
            },
            'fashion': {
                'fashion', 'style', 'outfit', 'clothing', 'apparel', 'designer',
                'trend', 'seasonal', 'collection', 'runway', 'boutique',
                'wardrobe', 'styling', 'lookbook', 'editorial'
            },
            'food': {
                'food', 'recipe', 'cooking', 'chef', 'restaurant', 'cuisine',
                'ingredients', 'meal', 'dining', 'foodie', 'culinary',
                'kitchen', 'gourmet', 'taste', 'flavor'
            },
            'tech': {
                'technology', 'software', 'app', 'digital', 'innovation',
                'platform', 'algorithm', 'data', 'analytics', 'AI',
                'machine learning', 'cloud', 'cybersecurity', 'blockchain'
            },
            'home': {
                'home', 'house', 'interior', 'decor', 'furniture', 'design',
                'renovation', 'DIY', 'garden', 'lifestyle', 'family',
                'domestic', 'household', 'living space'
            },
            'legal': {
                'contract', 'agreement', 'terms', 'conditions', 'liability',
                'indemnity', 'breach', 'remedy', 'jurisdiction', 'governing law',
                'dispute resolution', 'arbitration', 'confidentiality', 'proprietary'
            }
        }
        
        # Complexity indicators
        self.complexity_patterns = {
            'simple': {
                'patterns': [r'simple', r'basic', r'standard', r'regular'],
                'sentence_length_range': (10, 20),
                'clause_count_range': (1, 3)
            },
            'medium': {
                'patterns': [r'comprehensive', r'detailed', r'specific', r'include'],
                'sentence_length_range': (15, 30),
                'clause_count_range': (3, 6)
            },
            'complex': {
                'patterns': [r'sophisticated', r'advanced', r'extensive', r'comprehensive'],
                'sentence_length_range': (25, 50),
                'clause_count_range': (5, 10)
            }
        }
        
        # Clause pattern templates
        self.clause_templates = [
            r'the [A-Z][a-z]+ (?:shall|will|agrees to) .+',
            r'[A-Z][a-z]+ represents and warrants .+',
            r'in consideration of .+',
            r'subject to the terms .+',
            r'notwithstanding .+',
            r'for the purposes of .+'
        ]
    
    def load_all_templates(self) -> List[Dict[str, Any]]:
        """
        Load all templates from the templates directory.
        
        Returns:
            List of template data dictionaries
        """
        # Enhanced template directory search - look in multiple possible locations
        possible_dirs = [
            self.templates_dir,
            Path("./Thinkerbell_template_pipeline/thinkerbell/data"),
            Path("./thinkerbell/data"),
            Path("./data"),
        ]
        
        templates_dir = None
        for dir_path in possible_dirs:
            if dir_path.exists():
                # Check if it has PDF or DOCX files
                pdf_files = list(dir_path.glob("**/*.pdf"))
                docx_files = list(dir_path.glob("**/*.docx"))
                if pdf_files or docx_files:
                    templates_dir = dir_path
                    logger.info(f"Found template files in: {templates_dir}")
                    break
        
        if not templates_dir:
            logger.warning(f"No template directory found with PDF/DOCX files. Searched: {[str(d) for d in possible_dirs]}")
            return []

        templates = []
        
        # Find all supported files - focus on PDF and DOCX
        template_files = []
        template_files.extend(templates_dir.glob("**/*.pdf"))
        template_files.extend(templates_dir.glob("**/*.docx"))
        
        logger.info(f"Found {len(template_files)} template files in {templates_dir}")
        
        # Add progress bar for template processing
        template_progress = tqdm(template_files, desc="Processing templates", unit="file") if TQDM_AVAILABLE else template_files
        
        for template_file in template_progress:
            try:
                if TQDM_AVAILABLE:
                    template_progress.set_description(f"Processing {template_file.name[:30]}...")
                else:
                    logger.info(f"Processing template: {template_file.name}")
                
                # Extract raw text using enhanced extraction
                raw_text = self._extract_template_text_enhanced(template_file)
                
                if not raw_text or len(raw_text.strip()) < 100:
                    logger.warning(f"Template {template_file.name} has insufficient text content ({len(raw_text)} chars), skipping")
                    continue
                
                # Clean text
                clean_text = self.cleaner.clean_text(raw_text)
                
                # Create template data
                template_data = {
                    'file_path': str(template_file),
                    'file_name': template_file.name,
                    'raw_text': raw_text,
                    'clean_text': clean_text,
                    'char_count': len(clean_text),
                    'word_count': len(clean_text.split())
                }
                
                templates.append(template_data)
                logger.info(f"✓ Processed {template_file.name}: {len(clean_text)} chars, {len(clean_text.split())} words")
                
            except Exception as e:
                logger.error(f"Failed to process template {template_file}: {e}")
                continue
        
        self.templates = templates
        logger.info(f"Successfully loaded {len(templates)} templates")
        
        return templates
    
    def load_template(self, file_path: str) -> dict:
        """
        Load and parse a single template file (PDF, DOCX, or JSON)
        
        Args:
            file_path: Path to the template file
            
        Returns:
            dict: Parsed template data or None if failed
        """
        try:
            from datetime import datetime
            
            file_path = Path(file_path)
            
            if not file_path.exists():
                logger.error(f"Template file not found: {file_path}")
                return None
            
            # Extract raw text using enhanced extraction
            raw_text = self._extract_template_text_enhanced(file_path)
            
            if not raw_text or len(raw_text.strip()) < 100:
                logger.warning(f"Template {file_path.name} has insufficient text content ({len(raw_text)} chars)")
                return None
            
            # Clean text
            clean_text = self.cleaner.clean_text(raw_text)
            
            # Create template data structure
            template_data = {
                'file_path': str(file_path),
                'file_name': file_path.name,
                'raw_text': raw_text,
                'clean_text': clean_text,
                'content': clean_text,  # Alias for compatibility
                'file_type': file_path.suffix.lower().replace('.', ''),
                'file_size': file_path.stat().st_size,
                'char_count': len(clean_text),
                'word_count': len(clean_text.split()),
                'loaded_at': datetime.now().isoformat(),
                'metadata': self._extract_template_metadata(clean_text)
            }
            
            logger.info(f"✓ Loaded template {file_path.name}: {len(clean_text)} chars, {len(clean_text.split())} words")
            return template_data
            
        except Exception as e:
            logger.error(f"Failed to load template {file_path}: {e}")
            return None

    def _extract_template_metadata(self, content: str) -> dict:
        """Extract metadata from template content"""
        if not content:
            return {}
        
        content_lower = content.lower()
        
        # Basic metadata extraction
        metadata = {
            'has_fee_section': any(term in content_lower for term in ['fee', 'payment', 'cost', 'budget', 'compensation']),
            'has_deliverables': any(term in content_lower for term in ['deliverable', 'content', 'post', 'video', 'requirement']),
            'has_exclusivity': any(term in content_lower for term in ['exclusiv', 'restrict', 'non-compete', 'limitation']),
            'has_terms': any(term in content_lower for term in ['term', 'duration', 'period', 'time']),
            'has_intellectual_property': any(term in content_lower for term in ['intellectual property', 'copyright', 'trademark', 'rights']),
            'has_confidentiality': any(term in content_lower for term in ['confidential', 'non-disclosure', 'private', 'proprietary']),
            'complexity': self._determine_template_complexity(content),
            'word_count': len(content.split()) if content else 0,
            'estimated_industry': self._estimate_industry(content_lower)
        }
        
        return metadata
    
    def _determine_template_complexity(self, content: str) -> str:
        """Determine template complexity based on content"""
        if not content:
            return 'simple'
        
        word_count = len(content.split())
        
        if word_count < 500:
            return 'simple'
        elif word_count < 1500:
            return 'medium'
        else:
            return 'complex'
    
    def _estimate_industry(self, content_lower: str) -> str:
        """Estimate industry based on content keywords"""
        if not content_lower:
            return 'general'
        
        # Industry keyword matching
        for industry, keywords in self.industry_terms.items():
            if sum(1 for keyword in keywords if keyword in content_lower) >= 2:
                return industry
        
        return 'general'
    
    def _extract_template_text_enhanced(self, file_path: Path) -> str:
        """Enhanced text extraction from template files."""
        try:
            if file_path.suffix.lower() == '.pdf':
                return self._extract_pdf_enhanced(file_path)
            elif file_path.suffix.lower() == '.docx':
                return self._extract_docx_enhanced(file_path)
            else:
                return self.ingester._extract_text(file_path)
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {e}")
            return ""
    
    def _extract_pdf_enhanced(self, path: Path) -> str:
        """Enhanced PDF text extraction using PyMuPDF."""
        if not PYMUPDF_AVAILABLE:
            logger.warning("PyMuPDF not available, falling back to basic extraction")
            return self.ingester._extract_pdf(path)
            
        try:
            doc = fitz.open(str(path))
            text_parts = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Try to get structured text first
                text_dict = page.get_text("dict")
                page_text = ""
                
                for block in text_dict["blocks"]:
                    if "lines" in block:
                        for line in block["lines"]:
                            line_text = ""
                            for span in line["spans"]:
                                line_text += span["text"]
                            if line_text.strip():
                                page_text += line_text + "\n"
                
                # Fallback to simple text extraction if structured fails
                if not page_text.strip():
                    page_text = page.get_text()
                
                if page_text.strip():
                    text_parts.append(page_text)
            
            doc.close()
            result = '\n\n'.join(text_parts)
            logger.debug(f"Extracted {len(result)} characters from PDF {path.name}")
            return result
            
        except Exception as e:
            logger.error(f"Enhanced PDF extraction failed for {path}: {e}")
            # Fallback to basic extraction
            return self.ingester._extract_pdf(path)
    
    def _extract_docx_enhanced(self, path: Path) -> str:
        """Enhanced DOCX text extraction using python-docx."""
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not available, falling back to basic extraction")
            return self.ingester._extract_docx(path)
            
        try:
            doc = DocxDocument(str(path))
            text_parts = []
            
            # Extract paragraphs with better formatting
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
            
            # Extract tables with better structure
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text.strip():
                        text_parts.append(row_text)
            
            # Extract headers and footers if available
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            text_parts.append(f"HEADER: {text}")
                
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            text_parts.append(f"FOOTER: {text}")
            
            result = '\n\n'.join(text_parts)
            logger.debug(f"Extracted {len(result)} characters from DOCX {path.name}")
            return result
            
        except Exception as e:
            logger.error(f"Enhanced DOCX extraction failed for {path}: {e}")
            # Fallback to basic extraction
            return self.ingester._extract_docx(path)
    
    def build_style_profiles(self, group_by: str = 'auto') -> Dict[str, StyleProfile]:
        """
        Build style profiles from loaded templates.
        
        Args:
            group_by: How to group templates ('auto', 'industry', 'complexity', 'single')
            
        Returns:
            Dictionary mapping profile_id to StyleProfile
        """
        if not self.templates:
            self.load_all_templates()
        
        if not self.templates:
            logger.warning("No templates available for style profiling")
            return {}
        
        logger.info(f"Building style profiles from {len(self.templates)} templates")
        
        # Enhanced: Create one profile per template to ensure proper mapping
        profiles = {}
        
        # Add progress bar for profile building
        template_progress = tqdm(self.templates, desc="Building style profiles", unit="template") if TQDM_AVAILABLE else self.templates
        
        for template in template_progress:
            template_name = Path(template['file_name']).stem
            profile_id = f"profile_{template_name}"
            
            if TQDM_AVAILABLE:
                template_progress.set_description(f"Building profile: {template_name[:20]}...")
            else:
                logger.info(f"Building style profile for template: {template_name}")
            
            profile = StyleProfile(profile_id)
            profile.template_count = 1
            
            text = template['clean_text']
            
            # Enhanced style feature extraction
            self._analyze_tone_markers(profile, text)
            self._analyze_clause_patterns(profile, text)
            self._analyze_phrase_vocabulary(profile, text)
            self._analyze_sentence_structures(profile, text)
            self._analyze_industry_indicators(profile, text)
            self._calculate_formality_score(profile, text)
            self._analyze_complexity_indicators(profile, text)
            
            # NEW: Determine industry, complexity, and style keywords
            self._determine_industry(profile, text)
            self._determine_complexity(profile, text)
            self._extract_style_keywords(profile, text)
            
            profiles[profile_id] = profile
            
            # Log the profile details
            logger.info(f"✓ Created profile '{profile_id}': industry={profile.industry}, complexity={profile.complexity}, keywords={len(profile.style_keywords)}")
        
        self.style_profiles = profiles
        logger.info(f"Built {len(profiles)} style profiles")
        
        # Save mapping diagnostics to disk
        self.mapping_diagnostics.save_failures_to_disk()
        
        # Get diagnostics summary
        diagnostics_summary = self.mapping_diagnostics.get_failure_summary()
        logger.info(f"📊 Mapping diagnostics: {diagnostics_summary['total_failures']} failures recorded")
        logger.info(f"📊 Success rate: {diagnostics_summary['success_rate']:.1f}%")
        
        if diagnostics_summary['total_failures'] > 0:
            logger.warning(f"🔍 Mapping failures breakdown:")
            for failure_type, count in diagnostics_summary['failure_breakdown']['by_type'].items():
                logger.warning(f"  - {failure_type}: {count}")
        
        # Validation: Ensure no unknown mappings
        unknown_profiles = [p for p in profiles.values() if p.industry == 'unknown' or p.complexity == 'unknown']
        if unknown_profiles:
            logger.warning(f"Found {len(unknown_profiles)} profiles with unknown mappings")
            for profile in unknown_profiles:
                logger.warning(f"Profile {profile.profile_id}: industry={profile.industry}, complexity={profile.complexity}")
        
        if len(profiles) < len(self.templates):
            raise ValueError(f"Profile count ({len(profiles)}) is less than template count ({len(self.templates)})")
        
        return profiles
    
    def _group_templates(self, group_by: str) -> Dict[str, List[Dict[str, Any]]]:
        """Group templates by specified criteria."""
        
        if group_by == 'single':
            return {'unified': self.templates}
        
        if group_by == 'auto':
            # Auto-detect grouping based on template characteristics
            if len(self.templates) <= 3:
                return {'unified': self.templates}
            else:
                # Group by detected industry
                return self._group_by_detected_industry()
        
        if group_by == 'industry':
            return self._group_by_detected_industry()
        
        if group_by == 'complexity':
            return self._group_by_detected_complexity()
        
        # Default: single group
        return {'unified': self.templates}
    
    def _group_by_detected_industry(self) -> Dict[str, List[Dict[str, Any]]]:
        """Group templates by detected industry."""
        groups = defaultdict(list)
        
        for template in self.templates:
            text = template['clean_text'].lower()
            
            # Score each industry
            industry_scores = {}
            for industry, terms in self.industry_terms.items():
                score = sum(1 for term in terms if term in text)
                if score > 0:
                    industry_scores[industry] = score
            
            # Assign to highest scoring industry
            if industry_scores:
                best_industry = max(industry_scores.items(), key=lambda x: x[1])[0]
                groups[best_industry].append(template)
            else:
                groups['general'].append(template)
        
        return dict(groups)
    
    def _group_by_detected_complexity(self) -> Dict[str, List[Dict[str, Any]]]:
        """Group templates by detected complexity."""
        groups = defaultdict(list)
        
        for template in self.templates:
            text = template['clean_text']
            
            # Calculate complexity score
            avg_sentence_length = self._calculate_avg_sentence_length(text)
            clause_count = len(re.findall(r'[.!?]', text))
            
            if avg_sentence_length < 20 and clause_count < 10:
                complexity = 'simple'
            elif avg_sentence_length < 35 and clause_count < 20:
                complexity = 'medium'
            else:
                complexity = 'complex'
            
            groups[complexity].append(template)
        
        return dict(groups)
    
    def _analyze_tone_markers(self, profile: StyleProfile, text: str):
        """Analyze tone markers in text."""
        text_lower = text.lower()
        
        formal_count = sum(1 for marker in self.formal_indicators if marker in text_lower)
        casual_count = sum(1 for marker in self.casual_indicators if marker in text_lower)
        
        profile.tone_markers = {
            'formal_count': formal_count,
            'casual_count': casual_count,
            'formal_ratio': formal_count / (formal_count + casual_count + 1),
            'dominant_tone': 'formal' if formal_count > casual_count else 'casual'
        }
    
    def _analyze_clause_patterns(self, profile: StyleProfile, text: str):
        """Analyze common clause patterns."""
        patterns = []
        
        for template in self.clause_templates:
            matches = re.findall(template, text, re.IGNORECASE)
            if matches:
                patterns.extend(matches[:3])  # Limit to 3 examples per pattern
        
        profile.clause_patterns = patterns[:20]  # Limit total patterns
    
    def _analyze_phrase_vocabulary(self, profile: StyleProfile, text: str):
        """Analyze key phrase vocabulary."""
        # Extract significant phrases (2-4 words)
        words = re.findall(r'\b[A-Za-z]+\b', text.lower())
        
        # Find common bigrams and trigrams
        bigrams = [f"{words[i]} {words[i+1]}" for i in range(len(words)-1)]
        trigrams = [f"{words[i]} {words[i+1]} {words[i+2]}" for i in range(len(words)-2)]
        
        # Count frequencies
        bigram_counts = Counter(bigrams)
        trigram_counts = Counter(trigrams)
        
        # Keep most common phrases
        common_phrases = set()
        common_phrases.update([phrase for phrase, count in bigram_counts.most_common(20) if count > 1])
        common_phrases.update([phrase for phrase, count in trigram_counts.most_common(10) if count > 1])
        
        profile.phrase_vocabulary = common_phrases
    
    def _analyze_sentence_structures(self, profile: StyleProfile, text: str):
        """Analyze typical sentence structures."""
        sentences = re.split(r'[.!?]+', text)
        structures = []
        
        for sentence in sentences[:20]:  # Analyze first 20 sentences
            sentence = sentence.strip()
            if len(sentence) > 10:  # Skip very short sentences
                # Simplified structure analysis
                structure = {
                    'length': len(sentence.split()),
                    'has_subordinate_clause': ',' in sentence,
                    'starts_with_subject': self._starts_with_subject(sentence),
                    'contains_legal_terms': any(term in sentence.lower() for term in self.formal_indicators)
                }
                structures.append(structure)
        
        profile.sentence_structures = structures
    
    def _analyze_industry_indicators(self, profile: StyleProfile, text: str):
        """Analyze industry-specific indicators."""
        text_lower = text.lower()
        indicators = set()
        
        for industry, terms in self.industry_terms.items():
            found_terms = [term for term in terms if term in text_lower]
            if found_terms:
                indicators.update(found_terms)
        
        profile.industry_indicators = indicators
    
    def _calculate_formality_score(self, profile: StyleProfile, text: str):
        """Calculate overall formality score."""
        formal_count = profile.tone_markers.get('formal_count', 0)
        casual_count = profile.tone_markers.get('casual_count', 0)
        
        # Additional formality indicators
        avg_sentence_length = self._calculate_avg_sentence_length(text)
        passive_voice_count = len(re.findall(r'\b(?:is|are|was|were|being|been)\s+\w+ed\b', text))
        
        # Normalize scores
        formal_ratio = formal_count / (formal_count + casual_count + 1)
        length_score = min(avg_sentence_length / 30, 1.0)  # Cap at 30 words
        passive_score = min(passive_voice_count / 10, 1.0)  # Cap at 10 instances
        
        # Weighted average
        formality_score = (formal_ratio * 0.5 + length_score * 0.3 + passive_score * 0.2)
        profile.formality_score = formality_score
    
    def _analyze_complexity_indicators(self, profile: StyleProfile, text: str):
        """Analyze complexity indicators."""
        indicators = {}
        
        for complexity, data in self.complexity_patterns.items():
            pattern_count = 0
            for pattern in data['patterns']:
                pattern_count += len(re.findall(pattern, text, re.IGNORECASE))
            
            indicators[complexity] = {
                'pattern_count': pattern_count,
                'avg_sentence_length': self._calculate_avg_sentence_length(text),
                'clause_density': len(re.findall(r'[,;]', text)) / max(len(text.split()), 1)
            }
        
        profile.complexity_indicators = indicators
    
    def _calculate_avg_sentence_length(self, text: str) -> float:
        """Calculate average sentence length in words."""
        sentences = re.split(r'[.!?]+', text)
        lengths = [len(sentence.split()) for sentence in sentences if sentence.strip()]
        return sum(lengths) / len(lengths) if lengths else 0
    
    def _starts_with_subject(self, sentence: str) -> bool:
        """Check if sentence starts with a clear subject."""
        words = sentence.strip().split()
        if not words:
            return False
        
        first_word = words[0].lower()
        # Simple heuristic: starts with article, pronoun, or proper noun
        return first_word in ['the', 'a', 'an', 'this', 'that', 'i', 'you', 'we', 'they'] or words[0][0].isupper()
    
    def _determine_industry_and_complexity(self, template_data: Dict[str, Any]) -> Tuple[str, str]:
        """
        Determine industry and complexity from template content.
        
        ENHANCED: Better classification to avoid 'unknown' mappings.
        """
        text = template_data.get('clean_text', '').lower()
        
        # Industry determination with broader patterns
        industry_patterns = {
            'entertainment': [
                'talent', 'influencer', 'social media', 'content creator', 'youtube', 'instagram', 
                'tiktok', 'streaming', 'video', 'podcast', 'entertainment', 'celebrity',
                'brand ambassador', 'collaboration', 'campaign', 'promotion'
            ],
            'fashion': [
                'fashion', 'clothing', 'apparel', 'style', 'beauty', 'makeup', 'skincare',
                'accessories', 'jewelry', 'footwear', 'designer', 'boutique', 'retail'
            ],
            'food': [
                'food', 'restaurant', 'recipe', 'cooking', 'chef', 'culinary', 'dining',
                'beverage', 'nutrition', 'grocery', 'catering', 'cafe', 'bakery'
            ],
            'tech': [
                'technology', 'software', 'app', 'digital', 'platform', 'tech', 'startup',
                'innovation', 'ai', 'machine learning', 'data', 'analytics', 'saas'
            ],
            'health': [
                'health', 'wellness', 'fitness', 'medical', 'healthcare', 'therapy',
                'nutrition', 'supplement', 'exercise', 'yoga', 'mental health'
            ],
            'home': [
                'home', 'furniture', 'decor', 'interior', 'design', 'renovation',
                'appliance', 'garden', 'diy', 'cleaning', 'organization'
            ]
        }
        
        # Score each industry
        industry_scores = {}
        for industry, keywords in industry_patterns.items():
            score = sum(1 for keyword in keywords if keyword in text)
            if score > 0:
                industry_scores[industry] = score
        
        # Default to entertainment if no clear match (most agreements are influencer/talent)
        industry = 'entertainment'
        if industry_scores:
            industry = max(industry_scores, key=industry_scores.get)
        
        # Complexity determination based on multiple factors
        complexity_score = 0
        
        # Factor 1: Document length
        char_count = len(text)
        if char_count > 10000:
            complexity_score += 2
        elif char_count > 5000:
            complexity_score += 1
        
        # Factor 2: Legal terms complexity
        complex_terms = [
            'intellectual property', 'confidentiality', 'non-disclosure', 'termination',
            'indemnification', 'liability', 'jurisdiction', 'arbitration', 'breach',
            'governing law', 'force majeure', 'assignment', 'amendment'
        ]
        complexity_score += sum(1 for term in complex_terms if term in text)
        
        # Factor 3: Financial complexity
        financial_terms = ['fee', 'payment', 'royalty', 'commission', 'bonus', 'penalty']
        financial_count = sum(1 for term in financial_terms if term in text)
        if financial_count >= 3:
            complexity_score += 1
        
        # Factor 4: Number of sections/clauses (estimate by periods and line breaks)
        estimated_clauses = text.count('.') + text.count('\n')
        if estimated_clauses > 100:
            complexity_score += 1
        
        # Determine complexity
        if complexity_score >= 4:
            complexity = 'complex'
        elif complexity_score >= 2:
            complexity = 'medium'
        else:
            complexity = 'simple'
        
        # Log the determination for diagnostics
        self.mapping_diagnostics.record_successful_mapping(
            template_name=template_data.get('file_name', 'unknown'),
            industry=industry,
            complexity=complexity
        )
        
        return industry, complexity
    
    def _determine_industry(self, profile: StyleProfile, text: str):
        """Determine the industry for this profile based on text content."""
        text_lower = text.lower()
        industry_scores = {}
        
        # Score each industry based on term frequency
        for industry, terms in self.industry_terms.items():
            score = 0
            matched_terms = []
            for term in terms:
                count = text_lower.count(term.lower())
                if count > 0:
                    score += count
                    matched_terms.append(term)
            
            if score > 0:
                industry_scores[industry] = {
                    'score': score,
                    'matched_terms': matched_terms
                }
        
        # Determine best industry
        if industry_scores:
            best_industry = max(industry_scores.items(), key=lambda x: x[1]['score'])
            profile.industry = best_industry[0]
            
            # Check if this is a low-confidence match
            best_score = best_industry[1]['score']
            if best_score < 3:  # Low confidence threshold
                logger.warning(f"Low confidence industry mapping for {profile.profile_id}: '{profile.industry}' (score: {best_score})")
                
                # Record as a potential failure for analysis
                template_name = profile.profile_id.replace('profile_', '') if profile.profile_id.startswith('profile_') else profile.profile_id
                self.mapping_diagnostics.record_template_mapping_failure(
                    template_name=template_name,
                    template_text=text,
                    profile_id=profile.profile_id,
                    failed_attribute='industry',
                    attempted_value=profile.industry,
                    fallback_value=profile.industry,
                    scoring_data={'industry_scores': {k: v['score'] for k, v in industry_scores.items()}}
                )
            else:
                # Record successful mapping
                template_name = profile.profile_id.replace('profile_', '') if profile.profile_id.startswith('profile_') else profile.profile_id
                self.mapping_diagnostics.record_successful_mapping(
                    template_name=template_name,
                    profile_id=profile.profile_id,
                    industry=profile.industry
                )
            
            logger.debug(f"Determined industry '{profile.industry}' with score {best_score} for profile {profile.profile_id}")
        else:
            # Fallback: try to infer from file content
            fallback_industry = 'general'  # Default fallback
            
            if any(word in text_lower for word in ['talent', 'agreement', 'influencer', 'brand']):
                fallback_industry = 'entertainment'
            elif any(word in text_lower for word in ['contract', 'legal', 'terms', 'conditions']):
                fallback_industry = 'legal'
            
            profile.industry = fallback_industry
            
            # Record mapping failure - no industry terms found
            template_name = profile.profile_id.replace('profile_', '') if profile.profile_id.startswith('profile_') else profile.profile_id
            self.mapping_diagnostics.record_template_mapping_failure(
                template_name=template_name,
                template_text=text,
                profile_id=profile.profile_id,
                failed_attribute='industry',
                attempted_value=None,
                fallback_value=fallback_industry,
                scoring_data={'industry_scores': industry_scores}
            )
            
            logger.debug(f"Used fallback industry '{profile.industry}' for profile {profile.profile_id}")
    
    def _determine_complexity(self, profile: StyleProfile, text: str):
        """Determine the complexity level for this profile."""
        # Calculate various complexity metrics
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        if not sentences:
            profile.complexity = 'simple'
            
            # Record failure for empty/no sentences
            template_name = profile.profile_id.replace('profile_', '') if profile.profile_id.startswith('profile_') else profile.profile_id
            self.mapping_diagnostics.record_template_mapping_failure(
                template_name=template_name,
                template_text=text,
                profile_id=profile.profile_id,
                failed_attribute='complexity',
                attempted_value=None,
                fallback_value='simple',
                scoring_data={'complexity_metrics': {'sentence_count': 0, 'reason': 'no_sentences'}}
            )
            return
        
        # Average sentence length
        avg_sentence_length = sum(len(s.split()) for s in sentences) / len(sentences)
        
        # Clause density (commas and semicolons per sentence)
        clause_density = (text.count(',') + text.count(';')) / len(sentences)
        
        # Formal indicator density
        formal_count = sum(1 for indicator in self.formal_indicators if indicator in text.lower())
        formal_density = formal_count / len(text.split()) * 1000  # per 1000 words
        
        # Complexity scoring
        complexity_score = 0
        
        # Sentence length scoring
        if avg_sentence_length > 30:
            complexity_score += 3
        elif avg_sentence_length > 20:
            complexity_score += 2
        else:
            complexity_score += 1
        
        # Clause density scoring
        if clause_density > 0.3:
            complexity_score += 3
        elif clause_density > 0.15:
            complexity_score += 2
        else:
            complexity_score += 1
        
        # Formal density scoring
        if formal_density > 5:
            complexity_score += 3
        elif formal_density > 2:
            complexity_score += 2
        else:
            complexity_score += 1
        
        # Determine complexity level
        if complexity_score >= 8:
            profile.complexity = 'complex'
        elif complexity_score >= 5:
            profile.complexity = 'medium'
        else:
            profile.complexity = 'simple'
        
        # Prepare complexity metrics for diagnostics
        complexity_metrics = {
            'avg_sentence_length': avg_sentence_length,
            'clause_density': clause_density,
            'formal_density': formal_density,
            'complexity_score': complexity_score,
            'sentence_count': len(sentences)
        }
        
        # Check if complexity determination seems questionable
        template_name = profile.profile_id.replace('profile_', '') if profile.profile_id.startswith('profile_') else profile.profile_id
        
        # Record successful complexity mapping (most cases)
        if complexity_score >= 3:  # Reasonable confidence
            self.mapping_diagnostics.record_successful_mapping(
                template_name=template_name,
                profile_id=profile.profile_id,
                complexity=profile.complexity
            )
        else:
            # Record as potential failure - very low complexity score suggests poor text quality
            self.mapping_diagnostics.record_template_mapping_failure(
                template_name=template_name,
                template_text=text,
                profile_id=profile.profile_id,
                failed_attribute='complexity',
                attempted_value=profile.complexity,
                fallback_value=profile.complexity,
                scoring_data={'complexity_metrics': complexity_metrics}
            )
        
        logger.debug(f"Determined complexity '{profile.complexity}' (score: {complexity_score}) for profile {profile.profile_id}")
    
    def _extract_style_keywords(self, profile: StyleProfile, text: str):
        """Extract style keywords (top N frequent nouns/adjectives) from text."""
        keywords = []
        
        if self.nlp:
            # Use spaCy for better POS tagging
            try:
                doc = self.nlp(text[:100000])  # Limit text length for performance
                
                # Extract nouns and adjectives
                candidates = []
                for token in doc:
                    if (token.pos_ in ['NOUN', 'ADJ'] and 
                        len(token.text) > 3 and 
                        token.is_alpha and 
                        not token.is_stop and
                        not token.is_punct):
                        candidates.append(token.lemma_.lower())
                
                # Count frequencies
                keyword_counts = Counter(candidates)
                
                # Get top 15 most common
                keywords = [word for word, count in keyword_counts.most_common(15) if count > 1]
                
            except Exception as e:
                logger.warning(f"spaCy processing failed, using fallback: {e}")
                keywords = self._extract_keywords_fallback(text)
        else:
            # Fallback to basic keyword extraction
            keywords = self._extract_keywords_fallback(text)
        
        profile.style_keywords = keywords
        logger.debug(f"Extracted {len(keywords)} style keywords for profile {profile.profile_id}")
    
    def _extract_keywords_fallback(self, text: str) -> List[str]:
        """Fallback keyword extraction without spaCy."""
        # Simple approach: extract common words that might be nouns/adjectives
        words = re.findall(r'\b[A-Za-z]{4,}\b', text.lower())
        
        # Filter out common stop words and very common words
        stop_words = {
            'that', 'this', 'with', 'from', 'they', 'have', 'been', 'were', 'said',
            'each', 'which', 'their', 'time', 'will', 'about', 'would', 'there',
            'could', 'other', 'after', 'first', 'well', 'year', 'work', 'such',
            'make', 'even', 'most', 'some', 'take', 'than', 'only', 'think',
            'also', 'people', 'way', 'may', 'many', 'much', 'through', 'back',
            'good', 'very', 'still', 'should', 'any', 'our', 'out', 'up',
            'agreement', 'party', 'parties', 'shall', 'terms', 'conditions'
        }
        
        # Count word frequencies
        word_counts = Counter(word for word in words if word not in stop_words)
        
        # Return top 15 most common words
        return [word for word, count in word_counts.most_common(15) if count > 1]
    
    def save_style_profiles(self, output_dir: Path):
        """Save style profiles to JSON files."""
        ensure_dir(output_dir)
        
        # Save individual profiles
        for profile_id, profile in self.style_profiles.items():
            profile_file = output_dir / f"style_profile_{profile_id}.json"
            with open(profile_file, 'w', encoding='utf-8') as f:
                json.dump(profile.to_dict(), f, ensure_ascii=False, indent=2)
        
        # Save combined profiles index
        index_file = output_dir / "style_profiles_index.json"
        index_data = {
            'profiles': list(self.style_profiles.keys()),
            'profile_count': len(self.style_profiles),
            'template_count': sum(p.template_count for p in self.style_profiles.values())
        }
        
        with open(index_file, 'w', encoding='utf-8') as f:
            json.dump(index_data, f, ensure_ascii=False, indent=2)
        
        logger.info(f"Saved {len(self.style_profiles)} style profiles to {output_dir}")
    
    def load_style_profiles(self, input_dir: Path) -> Dict[str, StyleProfile]:
        """Load style profiles from JSON files."""
        profiles = {}
        
        # Load index
        index_file = input_dir / "style_profiles_index.json"
        if not index_file.exists():
            logger.warning(f"Style profiles index not found: {index_file}")
            return profiles
        
        with open(index_file, 'r', encoding='utf-8') as f:
            index_data = json.load(f)
        
        # Load each profile
        for profile_id in index_data['profiles']:
            profile_file = input_dir / f"style_profile_{profile_id}.json"
            if profile_file.exists():
                with open(profile_file, 'r', encoding='utf-8') as f:
                    profile_data = json.load(f)
                profiles[profile_id] = StyleProfile.from_dict(profile_data)
        
        self.style_profiles = profiles
        logger.info(f"Loaded {len(profiles)} style profiles from {input_dir}")
        
        return profiles
    
    def get_profile_summary(self) -> Dict[str, Any]:
        """Get summary of all style profiles."""
        if not self.style_profiles:
            return {}
        
        summary = {
            'profile_count': len(self.style_profiles),
            'total_templates': sum(p.template_count for p in self.style_profiles.values()),
            'profiles': {}
        }
        
        for profile_id, profile in self.style_profiles.items():
            summary['profiles'][profile_id] = {
                'template_count': profile.template_count,
                'formality_score': profile.formality_score,
                'dominant_tone': profile.tone_markers.get('dominant_tone', 'unknown'),
                'industry_indicators': list(profile.industry_indicators)[:5],  # Top 5
                'phrase_count': len(profile.phrase_vocabulary)
            }
        
        return summary 